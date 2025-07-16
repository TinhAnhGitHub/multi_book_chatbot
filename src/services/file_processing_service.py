import sys
import os
ROOT_DIR = os.path.abspath(
    os.path.join(os.path.dirname(__file__), '..')
)
sys.path.insert(0, ROOT_DIR)


from tqdm import tqdm
import mimetypes
import aiofiles
from pathlib import Path
from typing import cast
from datetime import datetime



from llama_index.core.schema import Document
from fastapi import UploadFile, HTTPException
import magic
import pymupdf4llm
import pymupdf
from docx2md import DocxFile,   DocxMedia, Converter
from docx import Document as DocxDocument

from src.utils.logger_config import SimpleLogger
from src.config.settings import AppConfig
from src.api.api_models import  FileValidation, SupportedFileType


logger = SimpleLogger(__name__)

class FileValidationService:
    """
    Service for file validation
    1. Validate file types
    """

    def __init__(self, config: AppConfig):
        self.config = config
        self.max_file_size = getattr(config, 'max_file_size', 50) * 1024 * 1024 # 50MB -> Bytes
        self.allowed_types = {
            SupportedFileType.PDF: ['application/pdf'],
            SupportedFileType.DOCX: [
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            ],
            SupportedFileType.TXT: ['text/plain'],
            SupportedFileType.MD: ['text/markdown', 'text/x-markdown']
        }

    
    async def validate_file(self, file: UploadFile) -> FileValidation:
        try:
            content = await file.read()
            await file.seek(0)

            file_size = len(content)
            if file_size > self.max_file_size:
                return FileValidation(
                    is_valid=False,
                    error_message=f"File size ({file_size / 1024 / 1024:.2f} MB) exceeds maximum allowed size ({self.max_file_size / 1024 / 1024} MB)"
                )
            

            if file_size == 0:
                return FileValidation(
                    is_valid=False,
                    error_message="File is empty"
                )
            
            try:
                detected_mime = magic.from_buffer(content, mime=True)
            except Exception as e:
                logger.warning(f"Could not detect mime type with magic: {e}")
                detected_mime = file.content_type
            
            file_type = self.get_filetype_from_mime(detected_mime)

            if not file_type:
                return FileValidation(
                    is_valid=False,
                    error_message=f"Unsupported file type: {detected_mime}. Supported types: {', '.join([ft.value for ft in SupportedFileType])}"
                )
            
            return FileValidation(
                is_valid=True,
                file_type=file_type.value,
                file_size=file_size
            )

        except Exception as e:
            logger.error(f"File validation error: {e}")
            return FileValidation(
                is_valid=False,
                error_message=f"Validation error: {str(e)}"
            )

    def get_filetype_from_mime(self, mime: str | None) -> SupportedFileType | None:
        if not mime:
            return None
    
        for file_type, mime_types in self.allowed_types.items():
            if mime in mime_types:
                return file_type
        return None
    



class FileStorageService:
    """
    Service for file storage management
    1. save uploaded file 
    2. delete files
    3. get_file_size
    """

    def __init__(self, config: AppConfig):
        self.config  = config
        self.storage_path = Path(config.data_path) / "uploads"

        self.storage_path.mkdir(parents=True, exist_ok=True)
    
    async def save_uploaded_file(self, file: UploadFile, file_id: str) -> str:

        filename = cast(str, file.filename)
        file_extension = Path(filename).suffix

        stored_filename = f"{file_id}{file_extension}"
        file_path = self.storage_path  / stored_filename

        async with aiofiles.open(file_path, 'wb') as f:
            content = await file.read()
            await f.write(content)
        
        return str(file_path)

    async def delete_file(self, file_path: str) -> bool:

        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                return True
            return False    
        except Exception as e:
            logger.error(f"Error deleting file {file_path}: {e}")
            return False
    
    async def get_file_size(self, file_path: str) -> int:
        try:
            return os.path.getsize(file_path)
        except Exception:
            return 0


class DocumentProcessingService:
    def __init__(self, config: AppConfig):
        self.config = config
    
    

    def from_pdf_to_document(self, file_path: str) -> Document:
        document = pymupdf.open(file_path)
        metadata_pdf = document.metadata

        metadata = {
            'producer': metadata_pdf.get('producer'),
            'author': metadata_pdf.get('author'),
            'keywords': metadata_pdf.get('keywords'),
            'title': metadata_pdf.get('title'),
            'subject': metadata_pdf.get('subject'),
            'file_name': os.path.basename(file_path)
        } if metadata_pdf else {}
        if not metadata_pdf:
                logger.warning(f"File {os.path.basename(file_path)} has no metadata.")


        markdown_content = pymupdf4llm.to_markdown(file_path)

        book_doc = Document(
            text=markdown_content,
            metadata=metadata
        )
        return book_doc
    
    def from_docx_to_document(self, file_path: str) -> Document:
        
        try:
            docx = DocxFile(file_path)
            converter = Converter(docx.document(), media=None, use_md_table=False)
            markdown_content = converter.convert()

            docx_meta = DocxDocument(file_path)
            core_props = docx_meta.core_properties

            metadata = {
                "title": core_props.title,
                "author": core_props.author,
                "subject": core_props.subject,
                "keywords": core_props.keywords,
                "last_modified_by": core_props.last_modified_by,
                "created": str(core_props.created),
                "modified": str(core_props.modified),
                "file_name": os.path.basename(file_path)
            }
            document = Document(
                text=markdown_content,
                metadata=metadata
            )

            return document

        except Exception as e:
            logger.error(f"Failed to process DOCX file {file_path}: {e}")
            raise
    
    def from_txt_md_to_document(self, file_path: str) -> Document:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            metadata = {
                "file_name": os.path.basename(file_path)
            }

            return Document(
                text=content, 
                metadata=metadata
            )

        except Exception as e:
            logger.error(f"Failed to process TXT file {file_path}: {e}")
            raise

    def from_file_to_document(self, file_path: str) -> Document:
        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".pdf":
            return self.from_pdf_to_document(file_path)
        elif ext == ".docx":
            return self.from_docx_to_document(file_path)
        elif ext == ".txt" or ext == '.md':
            return self.from_txt_md_to_document(file_path)
        else:
            raise ValueError(f"Unsupported file extension: {ext}")
        